"""
Word document processing utilities for filling templates with data
"""

import logging
from docx import Document
from docx.shared import Pt
from typing import Dict, Any, List

class WordProcessor:
    """Class to handle Word document template processing"""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
    
    def fill_template(self, template_path: str, data_list: List[Dict[str, Any]], output_path: str) -> bool:
        """
        Fill Word template with provided data from multiple rows
        
        Args:
            template_path: Path to the Word template file
            data_list: List of dictionaries containing the data to fill
            output_path: Path where the filled document will be saved
            
        Returns:
            True if successful, False otherwise
        """
        try:
            # Load the document
            doc = Document(template_path)
            
            # Find the first table in the document
            if not doc.tables:
                self.logger.error("No tables found in Word document")
                return False
            
            table = doc.tables[0]  # Use the first table
            
            # Ensure table has at least 2 rows
            if len(table.rows) < 2:
                self.logger.error("Table must have at least 2 rows")
                return False
            
            # Ensure table has at least 11 columns (updated for new columns)
            if len(table.columns) < 11:
                self.logger.error("Table must have at least 11 columns")
                return False
            
            # Process each row of data
            for i, data in enumerate(data_list):
                # Calculate which row to fill (starting from row 2, index 1)
                row_index = i + 1
                
                # Add more rows if needed
                while len(table.rows) <= row_index:
                    table.add_row()
                
                row = table.rows[row_index]
                
                # Map data to table columns with specific formatting and font sizes
                # Column 2 has 40 character limit
                numero_pedido = str(data.get('numero_pedido', '')).strip()[:40]
                
                self._fill_cell(row.cells[0], data.get('data'), 0)                    # Column 1 - Data (font 8)
                self._fill_cell(row.cells[1], numero_pedido, 1)                       # Column 2 - Número do Pedido (font 8, max 40 chars)
                self._fill_cell(row.cells[2], data.get('nome_cliente'), 2)            # Column 3 - Nome do Cliente (font 9)
                self._fill_cell(row.cells[3], data.get('prazo'), 3)                   # Column 4 - Prazo (font 8)
                self._fill_cell(row.cells[4], data.get('valor_pedido'), 4)            # Column 5 - Valor do Pedido (font 9)
                self._fill_cell(row.cells[5], data.get('porcentagem'), 5)             # Column 6 - Porcentagem (font 8)
                self._fill_cell(row.cells[6], data.get('valor_comissao'), 6)          # Column 7 - Valor da Comissão (font 9)
                self._fill_cell(row.cells[7], data.get('frete'), 7)                   # Column 8 - Frete (font 8)
                self._fill_cell(row.cells[8], data.get('referencia_comissao'), 8)     # Column 9 - Referência Comissão (font 9)
                self._fill_cell(row.cells[9], data.get('pagamento'), 9)               # Column 10 - Pagamento (font 8)
                
                # Column 11 can be left empty or filled with additional data if available
                if len(row.cells) > 10:
                    self._fill_cell(row.cells[10], "", 10)                            # Column 11 - Empty (font 8)
            
            # Save the filled document
            doc.save(output_path)
            
            self.logger.info(f"Successfully filled template with {len(data_list)} rows and saved to {output_path}")
            return True
            
        except Exception as e:
            self.logger.error(f"Error filling Word template: {str(e)}")
            return False
    
    def _fill_cell(self, cell, value: Any, column_index: int = 0) -> None:
        """
        Fill a table cell with the provided value and set font size
        
        Args:
            cell: The table cell object
            value: The value to insert
            column_index: Index of the column (0-based) for specific formatting
        """
        try:
            # Clear existing content
            cell.text = ""
            
            # Add new content
            if value is not None:
                # Format numeric values based on column
                if isinstance(value, (int, float)):
                    # Columns with no R$ symbol: 4,5,6,7,8 (valor_pedido, porcentagem, valor_comissao, frete, referencia_comissao)
                    if column_index in [4, 5, 6, 7, 8]:
                        if column_index == 5:  # Porcentagem column - format as integer
                            formatted_value = f"{int(value)}"
                        elif column_index == 7:  # Frete column - format as integer (no % symbol)
                            formatted_value = f"{int(value)}"
                        else:  # Other numeric columns - format with 2 decimals, no R$
                            formatted_value = f"{value:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
                        cell.text = formatted_value
                    else:
                        # Other numeric columns - keep default formatting
                        cell.text = str(value)
                else:
                    cell.text = str(value)
            else:
                cell.text = ""
            
            # Set font size based on column
            # Columns 1,2,4,6,8,10,11 (indexes 0,1,3,5,7,9,10) = font 8
            # Columns 3,5,7,9 (indexes 2,4,6,8) = font 9
            if column_index in [0, 1, 3, 5, 7, 9, 10]:  # Font size 8
                font_size = 8
            else:  # Font size 9 for columns 2,4,6,8 (indexes 2,4,6,8)
                font_size = 9
            
            # Apply font size to all paragraphs and runs in the cell
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)
                # If no runs exist, create one with the font size
                if not paragraph.runs and paragraph.text:
                    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(paragraph.text)
                    run.font.size = Pt(font_size)
                    
        except Exception as e:
            self.logger.warning(f"Error filling cell: {str(e)}")
            cell.text = str(value) if value is not None else ""
    
    def get_table_info(self, template_path: str) -> Dict[str, Any]:
        """
        Get information about tables in the Word document
        
        Args:
            template_path: Path to the Word template file
            
        Returns:
            Dictionary with table information
        """
        try:
            doc = Document(template_path)
            
            tables_info = []
            for i, table in enumerate(doc.tables):
                table_info = {
                    'index': i,
                    'rows': len(table.rows),
                    'columns': len(table.columns),
                    'headers': []
                }
                
                # Try to extract headers from first row
                if table.rows:
                    first_row = table.rows[0]
                    for cell in first_row.cells:
                        table_info['headers'].append(cell.text.strip())
                
                tables_info.append(table_info)
            
            return {
                'total_tables': len(doc.tables),
                'tables': tables_info,
                'has_suitable_table': any(
                    t['rows'] >= 2 and t['columns'] >= 7 
                    for t in tables_info
                )
            }
            
        except Exception as e:
            self.logger.error(f"Error getting table info: {str(e)}")
            return {'total_tables': 0, 'tables': [], 'has_suitable_table': False}
    
    def validate_template(self, template_path: str) -> bool:
        """
        Validate that the Word template has the required structure
        
        Args:
            template_path: Path to the Word template file
            
        Returns:
            True if template is valid, False otherwise
        """
        try:
            doc = Document(template_path)
            
            # Check if document has tables
            if not doc.tables:
                self.logger.error("Template must contain at least one table")
                return False
            
            # Check first table structure
            table = doc.tables[0]
            
            if len(table.rows) < 2:
                self.logger.error("Table must have at least 2 rows (header + data)")
                return False
            
            if len(table.columns) < 7:
                self.logger.error("Table must have at least 7 columns")
                return False
            
            return True
            
        except Exception as e:
            self.logger.error(f"Error validating template: {str(e)}")
            return False
    
    def create_sample_template(self, output_path: str) -> bool:
        """
        Create a sample Word template for testing purposes
        
        Args:
            output_path: Path where the sample template will be saved
            
        Returns:
            True if successful, False otherwise
        """
        try:
            doc = Document()
            
            # Add title
            title = doc.add_heading('Pedidos - Comissão 5% - mês', 0)
            
            # Add table
            table = doc.add_table(rows=2, cols=7)
            table.style = 'Table Grid'
            
            # Fill header row
            headers = [
                'Data', 'Nº Pedido', 'Nome do Cliente', 'Prazo p/ pagam.',
                'Valor Pedido', 'Porcentagem', 'Valor da Comissão'
            ]
            
            for i, header in enumerate(headers):
                table.cell(0, i).text = header
            
            # Add empty data row
            for i in range(7):
                table.cell(1, i).text = ""
            
            # Save document
            doc.save(output_path)
            
            self.logger.info(f"Sample template created at {output_path}")
            return True
            
        except Exception as e:
            self.logger.error(f"Error creating sample template: {str(e)}")
            return False
